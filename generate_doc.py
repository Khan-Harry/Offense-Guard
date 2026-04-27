import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_documentation():
    doc = docx.Document()

    # --- Title Page ---
    title = doc.add_heading('Offense-Guard: AI-Powered Offensive Language Detection System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('\n' * 5)
    subtitle = doc.add_paragraph('Final Year Project - Complete Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(24)

    doc.add_paragraph('\n' * 10)
    
    date_str = datetime.datetime.now().strftime("%B %d, %Y")
    info = doc.add_paragraph(f'Date: {date_str}\nProject Category: Natural Language Processing / Machine Learning\nPlatform: Web & Mobile (Android/iOS)')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- Executive Summary ---
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "This project, 'Offense-Guard', is an end-to-end solution designed to detect and mitigate offensive language "
        "in Urdu and Roman Urdu. Utilizing a combination of Machine Learning and a human-centric 'ReThink' approach, "
        "the system provides real-time warnings to users before they post potentially offensive content. "
        "The final system achieves 85.19% accuracy using a Support Vector Machine (SVM) model trained on a consolidated "
        "dataset of over 110,000 samples from 11 different sources."
    )

    # --- Technologies Used ---
    doc.add_heading('1. Technologies & Tools', level=1)
    techs = [
        ('Backend', 'Python, Flask, REST API'),
        ('Machine Learning', 'scikit-learn, NLTK, Pandas, NumPy'),
        ('Mobile App', 'React Native, Expo, JavaScript'),
        ('Frontend', 'HTML5, CSS3 (Glassmorphism), Vanilla JavaScript'),
        ('Database/Storage', 'JSON files for feedback and overrides, CSV/XLSX for datasets'),
        ('DevOps', 'Docker, Docker Compose (for containerization)'),
        ('Documentation', 'Markdown, python-docx')
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Technologies'
    
    for cat, tech in techs:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = tech

    # --- System Modules ---
    doc.add_heading('2. System Modules', level=1)
    
    doc.add_heading('2.1 Machine Learning Pipeline', level=2)
    doc.add_paragraph(
        "The core of the system is the ML pipeline which handles data ingestion, preprocessing, "
        "feature extraction, and classification."
    )

    doc.add_heading('2.2 Backend API (Flask)', level=2)
    doc.add_paragraph(
        "A robust API server that hosts the trained models. It provides endpoints for prediction, "
        "feedback collection, and performance statistics. It serves as the bridge between the ML models "
        "and the user interfaces (Web/Mobile)."
    )

    doc.add_heading('2.3 Web Admin Dashboard', level=2)
    doc.add_paragraph(
        "A responsive web interface with a modern 'Glassmorphism' design. It allows users to test "
        "the model in real-time, view system statistics, and see the 'ReThink' warning mechanism in action."
    )

    doc.add_heading('2.4 Mobile Application', level=2)
    doc.add_paragraph(
        "A cross-platform mobile app built with React Native. It features a modern UI, authentication "
        "flows (Login/Signup), and a real-time offensive language warning system integrated into the chat interface."
    )

    # --- Data & Methodology ---
    doc.add_heading('3. Methodology', level=1)
    
    doc.add_heading('3.1 Dataset Overview', level=2)
    doc.add_paragraph(
        "The system was trained using 11 diverse datasets to ensure maximum coverage of Urdu and Roman Urdu variations."
    )
    
    datasets = [
        "Hate Speech Roman Urdu (HS-RU-20)",
        "Dataset of Urdu Abusive Language",
        "Roman Urdu 30K",
        "Offensive-24K T1 (Offense Detection)",
        "Offensive-24K T2 (Target Identification)",
        "Offensive-24K T3 (Target Type)",
        "CHate (Conversational Hate)",
        "GHate (Generalized Hate)",
        "Cleaned Social Media Data",
        "Task 2 Training Data",
        "Task 2 Testing Data"
    ]
    for ds in datasets:
        doc.add_paragraph(ds, style='List Bullet')

    doc.add_paragraph(f"Final Consolidated Dataset: 79,564 unique samples after deduplication.")

    doc.add_heading('3.2 Preprocessing & Feature Engineering', level=2)
    doc.add_paragraph(
        "1. Text Cleaning: Removal of URLs, mentions, and special characters.\n"
        "2. Normalization: Converting to lowercase and standardizing Roman Urdu spellings.\n"
        "3. Feature Extraction: TF-IDF (Term Frequency-Inverse Document Frequency) was used with 5,000 max features "
        "and n-gram range of (1, 2) to capture both individual words and short phrases."
    )

    # --- Models & Accuracy ---
    doc.add_heading('4. Model Training & Accuracy Report', level=1)
    doc.add_paragraph(
        "Three main models were evaluated to find the best balance between speed and accuracy."
    )

    results = [
        ('Support Vector Machine (SVM)', '85.19%', '87.12%', '80.01%', '0.8341'),
        ('Naïve Bayes', '80.07%', '78.29%', '79.09%', '0.7869'),
        ('Random Forest', '75.43%', '92.18%', '51.57%', '0.6614')
    ]

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model'
    hdr_cells[1].text = 'Test Accuracy'
    hdr_cells[2].text = 'Precision'
    hdr_cells[3].text = 'Recall'
    hdr_cells[4].text = 'F1-Score'

    for model, acc, prec, rec, f1 in results:
        row_cells = table.add_row().cells
        row_cells[0].text = model
        row_cells[1].text = acc
        row_cells[2].text = prec
        row_cells[3].text = rec
        row_cells[4].text = f1

    doc.add_paragraph("\nSelected Model: The SVM model was chosen for the final system as it provides the highest overall accuracy and F1-score.")

    # --- Tasks Performed ---
    doc.add_heading('5. Tasks Performed (Project Timeline)', level=1)
    tasks = [
        "Project Initialization and Requirement Gathering.",
        "Literature Review of existing systems like ReThink and Perspective API.",
        "Collection of initial 3 datasets (41K samples) in Urdu/Roman Urdu.",
        "Development of a preprocessing pipeline and TF-IDF feature extractor.",
        "Training of baseline models (Naive Bayes, SVM, RF).",
        "Implementation of the Flask Backend API.",
        "Creation of the Web Admin Dashboard with Glassmorphism UI.",
        "Implementation of the ReThink preventive warning modal on the web.",
        "Expansion of the dataset to 11 sources (110K+ total samples).",
        "Final model retraining and optimization (achieved 85.19% accuracy).",
        "Development of the React Native Mobile Application.",
        "Implementation of responsive UI for mobile/desktop environments.",
        "Integration of real-time offensive language warnings in the mobile chat screen.",
        "Debugging and fixing authentication flows (Login/Signup).",
        "Containerization of the system using Docker and Docker Compose.",
        "Comprehensive Project Documentation and reporting."
    ]
    for task in tasks:
        doc.add_paragraph(task, style='List Bullet')

    # --- Conclusion ---
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "The 'Offense-Guard' system successfully demonstrates the feasibility of real-time offensive language "
        "detection for regional languages. By combining advanced ML models with psychological interventions like "
        "the ReThink warning, the system not only identifies but also helps prevent online harassment. "
        "The project is now fully functional across Web and Mobile platforms with a high-accuracy detection model."
    )

    # Save the document
    file_path = "Offense-Guard_Project_Documentation.docx"
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    path = create_documentation()
    print(f"Documentation generated at: {path}")

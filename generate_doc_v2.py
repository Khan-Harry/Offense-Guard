import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    for edge in ('top', 'start', 'bottom', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, value in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(value))

def create_documentation():
    doc = docx.Document()

    # --- Title Page ---
    title = doc.add_heading('Offense-Guard: Intelligent Offensive Language Detection System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('\n' * 3)
    subtitle = doc.add_paragraph('Comprehensive Final Year Project Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(28)
    subtitle.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph('\n' * 2)
    abstract_para = doc.add_paragraph(
        "A multi-tier AI system for real-time detection of offensive language in Urdu and Roman Urdu, "
        "featuring a Flask backend with MongoDB, high-performance ML/DL models, and a cross-platform mobile application."
    )
    abstract_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('\n' * 5)
    
    date_str = datetime.datetime.now().strftime("%B %d, %Y")
    info = doc.add_paragraph(f'Date: {date_str}\nVersion: 2.0 (Verified System Analysis)')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # --- Table of Contents Placeholder ---
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "Offense-Guard is a sophisticated AI-powered system designed to address the challenge of cyberbullying and "
        "offensive content in regional languages (Urdu and Roman Urdu). Unlike traditional keyword-based filters, "
        "this system utilizes advanced Natural Language Processing (NLP) and Machine Learning techniques to understand "
        "context, slang, and cultural nuances."
    )

    # --- Tools & Technologies ---
    doc.add_heading('2. Tools & Technologies', level=1)
    
    doc.add_heading('2.1 Backend & Database', level=2)
    doc.add_paragraph(
        "• Flask (Python): Used to build the RESTful API for model serving.\n"
        "• MongoDB: NoSQL database for managing users, feedback, and prediction history.\n"
        "• JWT (JSON Web Token): Secure authentication mechanism for API access.\n"
        "• PyMongo: Driver for database interactions."
    )

    doc.add_heading('2.2 Machine Learning & AI', level=2)
    doc.add_paragraph(
        "• Scikit-learn: Implementation of Traditional ML models (SVM, Naive Bayes, Random Forest).\n"
        "• TensorFlow/Keras: Implementation of Deep Learning architectures (CNN, LSTM).\n"
        "• NLTK & Pandas: Text preprocessing and data manipulation.\n"
        "• Word2Vec: Used for generating dense vector embeddings for DL models."
    )

    doc.add_heading('2.3 Mobile & Frontend', level=2)
    doc.add_paragraph(
        "• React Native: Framework for building the cross-platform mobile application.\n"
        "• Expo: Tooling for rapid mobile development and deployment.\n"
        "• CSS Glassmorphism: Modern UI design for the web admin dashboard."
    )

    doc.add_heading('2.4 DevOps', level=2)
    doc.add_paragraph(
        "• Docker: Containerization of the backend services.\n"
        "• Docker Compose: Orchestration of the Flask app and MongoDB services."
    )

    # --- System Architecture ---
    doc.add_heading('3. System Modules', level=1)
    
    doc.add_heading('3.1 Multi-Model Pipeline', level=2)
    doc.add_paragraph(
        "The system employs a hierarchical detection strategy:\n"
        "1. Manual Overrides: Immediate lookup in 'overrides.json' for whitelisted/blacklisted terms.\n"
        "2. SVM Classifier: The primary ML model using TF-IDF features for high-precision detection.\n"
        "3. LSTM/CNN: Experimental deep learning models used for context-aware verification when the primary model is uncertain."
    )

    doc.add_heading('3.2 Unified Data Loader', level=2)
    doc.add_paragraph(
        "A custom 'DataLoader' module maps 11 diverse datasets into a unified 4-class classification system:\n"
        "• Category 0: Neutral (Safe)\n"
        "• Category 1: Hate Speech\n"
        "• Category 2: Abusive/Profanity\n"
        "• Category 3: Offensive"
    )

    doc.add_heading('3.3 Mobile Integration', level=2)
    doc.add_paragraph(
        "The mobile app features a real-time 'ReThink' warning system. When a user types a message, "
        "it is sent to the `/predict` endpoint. If classified as offensive, the app triggers a preventive modal "
        "asking the user to reconsider before posting."
    )

    # --- Data & Accuracy ---
    doc.add_heading('4. Model Performance & Datasets', level=1)
    
    doc.add_heading('4.1 Dataset Statistics', level=2)
    doc.add_paragraph(
        "The system was trained on 110,000+ raw samples, resulting in 79,564 unique samples after "
        "rigorous cleaning and deduplication. Datasets include:\n"
        "• HS-RU-20, Urdu Abusive Language, Roman Urdu 30K, CHate, GHate, and Task-specific datasets."
    )

    doc.add_heading('4.2 Accuracy Report', level=2)
    
    results = [
        ('SVM (Primary)', '85.19%', '0.8341', 'Primary Model'),
        ('CNN (Deep Learning)', '82.45%', '0.8012', 'Secondary/Experimental'),
        ('LSTM (Deep Learning)', '83.12%', '0.8150', 'Secondary/Experimental'),
        ('Naive Bayes', '80.07%', '0.7869', 'Baseline'),
        ('Random Forest', '75.43%', '0.6614', 'Baseline')
    ]

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Model Architecture'
    hdr_cells[1].text = 'Accuracy'
    hdr_cells[2].text = 'F1-Score'
    hdr_cells[3].text = 'System Role'

    for model, acc, f1, role in results:
        row_cells = table.add_row().cells
        row_cells[0].text = model
        row_cells[1].text = acc
        row_cells[2].text = f1
        row_cells[3].text = role

    # --- Implementation Details ---
    doc.add_heading('5. Core Implementation Details', level=1)
    
    doc.add_heading('5.1 Backend Endpoints', level=2)
    doc.add_paragraph(
        "• POST /predict: Analyzes text and returns class (Neutral, Offensive, etc.) with confidence score.\n"
        "• POST /api/auth/login: Authenticates users and returns a JWT token.\n"
        "• GET /stats: Returns usage statistics and history for the authenticated user.\n"
        "• POST /feedback: Allows users to submit corrections for model training (Active Learning)."
    )

    doc.add_heading('5.2 Preprocessing Logic', level=2)
    doc.add_paragraph(
        "The system handles Roman Urdu spelling variations using a normalization dictionary (e.g., 'kia' -> 'kya', 'hy' -> 'hai'). "
        "It also performs character-level cleaning, URL removal, and whitespace normalization."
    )

    # --- Task Summary ---
    doc.add_heading('6. Tasks Performed (Comprehensive List)', level=1)
    tasks = [
        "Analyzed 11 datasets for Urdu/Roman Urdu offensive language.",
        "Implemented a robust DataLoader with multi-class mapping logic.",
        "Developed a Flask REST API with MongoDB integration for persistent storage.",
        "Implemented JWT-based secure authentication for mobile/web users.",
        "Trained and evaluated SVM, Naive Bayes, and Random Forest models.",
        "Architected and trained Deep Learning models (CNN and LSTM) using Keras.",
        "Developed a cross-platform Mobile Application using React Native/Expo.",
        "Created a Web Admin Dashboard with modern UI components.",
        "Implemented a real-time 'ReThink' warning mechanism across platforms.",
        "Set up an Active Learning feedback loop for incremental model improvement.",
        "Configured Docker/Docker Compose for seamless system deployment.",
        "Integrated prediction history and usage statistics for personalized user experience."
    ]
    for task in tasks:
        doc.add_paragraph(task, style='List Bullet')

    # --- Conclusion ---
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        "Offense-Guard successfully bridges the gap between AI research and practical application. "
        "By providing a high-accuracy, multi-platform solution for regional language moderation, "
        "it offers a tangible tool for creating safer digital spaces. The modular architecture ensures "
        "that as language evolves, the system can be easily updated through its integrated feedback mechanism."
    )

    # Save the document
    file_path = "Offense-Guard_Project_Documentation_V2.docx"
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    path = create_documentation()
    print(f"Documentation generated at: {path}")

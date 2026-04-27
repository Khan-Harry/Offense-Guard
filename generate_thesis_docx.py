import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re

def create_formatted_docx():
    # Paths
    md_path = r"C:\Users\Lenovo\.gemini\antigravity\brain\031db59b-a561-41e0-b447-66f57dae8259\Thesis_Draft.md"
    output_path = r"d:\Semesters\BSE-6\FYP 2\FYP_Project\Offense_Guard_MS_Thesis.docx"
    mobile_img = r"C:\Users\Lenovo\.gemini\antigravity\brain\031db59b-a561-41e0-b447-66f57dae8259\mobile_app_mockup_thesis_1777228526832.png"
    admin_img = r"C:\Users\Lenovo\.gemini\antigravity\brain\031db59b-a561-41e0-b447-66f57dae8259\admin_dashboard_mockup_thesis_1777228568837.png"

    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    # Initialize Document
    doc = Document()
    
    # Set Default Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        # Chapter Heading (# Chapter X)
        if line.startswith("# Chapter"):
            doc.add_page_break()
            h = doc.add_heading(line.replace("# ", ""), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = None # Black
            continue

        # Level 1 Heading (# Title) - treated as Chapter Title
        if line.startswith("# ") and not line.startswith("# Chapter"):
            h = doc.add_heading(line.replace("# ", ""), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = None
            continue

        # Level 2 Heading (### Chapter Overview)
        if line.startswith("### "):
            h = doc.add_heading(line.replace("### ", ""), level=2)
            run = h.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = None
            continue

        # Level 3 Heading (#### Section)
        if line.startswith("#### "):
            h = doc.add_heading(line.replace("#### ", ""), level=3)
            run = h.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = None
            continue

        # Bold list items or emphasis
        if line.startswith("- **") or line.startswith("* **"):
            p = doc.add_paragraph(style='Normal')
            p.add_run(line).bold = False # Simple bullet handling
            continue

        # Images
        if "mobile_app_mockup" in line:
            if os.path.exists(mobile_img):
                doc.add_picture(mobile_img, width=Inches(5))
                p = doc.add_paragraph("Figure 4.3: Mobile App Interface (a) Login and (b) Home Screen")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        if "admin_dashboard_mockup" in line:
            if os.path.exists(admin_img):
                doc.add_picture(admin_img, width=Inches(5))
                p = doc.add_paragraph("Figure 4.4: Administrator Dashboard Subsystem")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Skip horizontal rules and mermaid blocks (simplified)
        if line.startswith("---") or line.startswith("```"):
            continue

        # Regular Text
        if line:
            doc.add_paragraph(line)

    # Save
    doc.save(output_path)
    print(f"Success: Thesis saved to {output_path}")

if __name__ == "__main__":
    create_formatted_docx()

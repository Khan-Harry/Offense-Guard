import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('DevOps Assignment: Automating FYP Development Using GitHub Actions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph(
        "For this assignment, the Final Year Project (FYP) titled 'Offensive Language Detection App In Urdu' "
        "has been utilized. The project consists of a Python-based backend API (Flask) and a mobile frontend. "
        "A proper project folder structure has been established with source code, tests, and documentation."
    )
    
    # Workflow
    doc.add_heading('2. GitHub Actions Workflow Configuration', level=1)
    doc.add_paragraph(
        "A GitHub Actions CI/CD workflow has been implemented to automate the testing process. "
        "The workflow is triggered on both 'push' and 'pull_request' events to the main branches."
    )
    
    doc.add_paragraph("The workflow performs the following tasks:", style='List Bullet')
    doc.add_paragraph("Checks out the repository code.", style='List Bullet')
    doc.add_paragraph("Sets up the Python environment (Python 3.10).", style='List Bullet')
    doc.add_paragraph("Installs project dependencies and required testing tools (pytest, flake8).", style='List Bullet')
    doc.add_paragraph("Runs code quality and linting checks using flake8.", style='List Bullet')
    doc.add_paragraph("Executes the test suite automatically using pytest to verify code integrity.", style='List Bullet')
    
    doc.add_paragraph("Below is the content of the workflow YAML file (.github/workflows/ci.yml):")
    
    # Code block
    code = '''name: Continuous Integration

on:
  push:
    branches: [ "main", "master" ]
  pull_request:
    branches: [ "main", "master" ]

jobs:
  test:
    runs-on: ubuntu-latest
    steps:
    - name: Checkout Repository
      uses: actions/checkout@v3
    - name: Set up Python
      uses: actions/setup-python@v4
      with:
        python-version: '3.10'
    - name: Install dependencies
      run: |
        python -m pip install --upgrade pip
        pip install pytest flake8
        if [ -f requirements.txt ]; then pip install -r requirements.txt; fi
    - name: Lint with flake8
      run: |
        flake8 . --count --select=E9,F63,F7,F82 --show-source --statistics
        flake8 . --count --exit-zero --max-complexity=10 --max-line-length=127 --statistics
    - name: Run Tests with pytest
      run: |
        pytest tests/'''
    
    code_para = doc.add_paragraph(code)
    code_para.style.font.name = 'Courier New'
    code_para.style.font.size = Pt(9)
    
    # Page Break
    doc.add_page_break()
    
    # Screenshots Section
    doc.add_heading('3. Evidence of Successful Execution', level=1)
    doc.add_paragraph(
        "This section provides visual evidence of the GitHub Actions workflow executing successfully "
        "after a code push."
    )
    
    # Placeholders for screenshots
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ INSERT SCREENSHOT OF GITHUB ACTIONS WORKFLOW RUN HERE ]")
    run.bold = True
    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
    
    doc.add_paragraph("\n\n")
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("[ INSERT SCREENSHOT OF PASSING TESTS/LOGS HERE ]")
    run2.bold = True
    run2.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
    
    doc.add_heading('4. Conclusion', level=1)
    doc.add_paragraph(
        "By integrating GitHub Actions, the development pipeline for the Offensive Language Detection project "
        "has been significantly enhanced. Automated testing ensures that any new code changes are immediately "
        "verified against the test suite, preventing bugs from reaching the main branch and ensuring "
        "higher software quality and reliability."
    )
    
    # Save document
    doc.save('DevOps_Assignment_01_Report.docx')
    print("Report generated successfully as 'DevOps_Assignment_01_Report.docx'")

if __name__ == '__main__':
    create_report()
